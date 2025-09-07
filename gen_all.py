#!/usr/bin/env -S uv run
# /// script
# requires-python = ">=3.13"
# dependencies = [
#   "pandas>=2.3.2",
#   "openpyxl>=3.1.5",
#   "python-docx>=1.1.2",
#   "rich>=14.1.0",
# ]
# ///

from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from rich.console import Console
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


WORKBOOK_NAME = "rosters.xlsx"
TEACHER_CSV_NAME = "room_teachers.csv"
OUTPUT_DIR = Path("out")


console = Console()


def load_workbook(path: Path) -> Dict[str, pd.DataFrame]:
    if not path.exists():
        raise FileNotFoundError(
            f"Workbook not found: {path}. Place '{WORKBOOK_NAME}' in the project directory."
        )
    frames: Dict[str, pd.DataFrame] = pd.read_excel(
        path, sheet_name=None, engine="openpyxl", dtype=object
    )
    return frames


def load_teacher_mapping(csv_path: Path) -> Dict[str, str]:
    if not csv_path.exists():
        console.print(
            f"[yellow]Warning:[/] teacher map CSV not found at {csv_path}. The Teacher field will be blank."
        )
        return {}
    df = pd.read_csv(csv_path)
    mapping: Dict[str, str] = {}
    for _, row in df.iterrows():
        room = str(row.get("room", "")).strip()
        teacher = str(row.get("teacher", "")).strip()
        if room:
            mapping[room] = teacher
    return mapping


# ----------------------- Shared helpers ----------------------- #

def parse_homeroom_from_header(df: pd.DataFrame) -> Optional[int]:
    for col in df.columns:
        if isinstance(col, str) and col.strip().lower().startswith("homeroom "):
            suffix = col.strip().split(" ", 1)[1]
            try:
                return int(str(suffix).strip())
            except ValueError:
                continue
    if "Room" in df.columns:
        for value in df["Room"].tolist():
            if pd.notna(value):
                try:
                    return int(str(value).strip())
                except Exception:
                    pass
    return None


def extract_names_for_sheet(df: pd.DataFrame) -> List[str]:
    name_column: Optional[str] = None
    first_name_column: Optional[str] = None

    for col in df.columns:
        if isinstance(col, str) and col.strip().lower().startswith("homeroom "):
            name_column = col
            break

    for col in df.columns:
        if isinstance(col, str) and (col.strip().lower() in {"first", "first name"} or col.startswith("Unnamed")):
            first_name_column = col
            break

    names: List[str] = []
    if name_column is not None and first_name_column is not None:
        last_series = df[name_column]
        first_series = df[first_name_column]
        for last, first in zip(last_series.tolist(), first_series.tolist()):
            if pd.isna(last) and pd.isna(first):
                continue
            last_part = str(last).strip() if pd.notna(last) else ""
            first_part = str(first).strip() if pd.notna(first) else ""
            full = f"{last_part}, {first_part}".strip().strip(",")
            if full:
                names.append(full)
        return names

    if name_column is not None:
        for value in df[name_column].tolist():
            if pd.isna(value):
                continue
            names.append(str(value).strip())
        return names

    for col in df.columns:
        series = df[col]
        if series.dtype == object:
            for value in series.tolist():
                if pd.notna(value):
                    names.append(str(value).strip())
            if names:
                return names
    return names


def grade_label_for_sheet(sheet_name: str) -> str:
    lowered = sheet_name.strip().lower()
    if lowered in {"k", "kinder", "kindergarten"}:
        return "Kinder"
    if lowered in {"prek", "pre-k", "pre k", "new prek", "new pre-k"}:
        return "PreK"
    return sheet_name.strip()


# ----------------------- Card generator (PreK/K/1st) ----------------------- #

def add_student_card(document: Document, student_name: str, teacher: str, room: Optional[int], grade: str) -> None:
    title = document.add_paragraph()
    run_label = title.add_run("Student Name:  ")
    run_label.bold = True
    run_label.font.size = Pt(24)
    run_name = title.add_run(student_name)
    run_name.bold = True
    run_name.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Teacher", "Homeroom", "Grade"]
    values = [teacher or "", str(room) if room is not None else "", grade]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
    for idx, text in enumerate(values):
        cell = table.rows[1].cells[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
    document.add_page_break()


def generate_cards(frames: Dict[str, pd.DataFrame], teacher_by_room: Dict[str, str], outdir: Path) -> None:
    for sheet in ["New PreK", "kinder", "1st"]:
        if sheet not in frames:
            console.print(f"[yellow]Skipping missing sheet:[/] {sheet}")
            continue
        df = frames[sheet]
        names = extract_names_for_sheet(df)
        room_num = parse_homeroom_from_header(df)
        grade = grade_label_for_sheet(sheet)
        teacher = teacher_by_room.get(str(room_num), "")
        document = Document()
        for name in names:
            add_student_card(document, name, teacher, room_num, grade)
        out_path = outdir / (sheet.replace("/", "-").replace(" ", "_") + ".docx")
        document.save(out_path)
        console.print(f"[green]Wrote[/] {out_path}")


# ----------------------- 2nd–5th schedules ----------------------- #

SCHEDULE_ROWS_2_5: List[tuple[str, str]] = [
    ("8:00-8:10", "Homeroom/Breakfast/Calm Classroom"),
    ("8:10-9:10", "Reading/Writing-Class 1"),
    ("9:10-10:10", "Student Enrichment"),
    ("10:10-10:40", "Reading/Writing-Class 1"),
    ("10:40-11:30", "SS-Class 1"),
    ("11:30-12:00", "Class 1:\n*Second Step on Monday and Friday\n*MTSS on Tuesday, Wednesday Thursday"),
    ("12:00-12:30", "Recess-Indoor/Lunchroom A"),
    ("12:30-12:55", "Lunch-Lunchroom B"),
    ("12:55-2:25", "Calm Classroom/Math-Class 2"),
    ("2:25-2:55", "Science-Class 2"),
    ("2:55-3:00", "Homeroom/Dismissal"),
]


def add_header(document: Document, name: str, homeroom: Optional[int], grade: str, language: Optional[str]) -> None:
    line1 = document.add_paragraph()
    left = line1.add_run(f"Name:{name}")
    left.bold = True
    left.font.size = Pt(18)
    line1.add_run("\t\t\t")
    right = line1.add_run(f"Grade:{grade}")
    right.bold = True
    right.font.size = Pt(18)

    line2 = document.add_paragraph()
    l2a = line2.add_run(f"Homeroom: {homeroom if homeroom is not None else ''}")
    l2a.bold = True
    l2a.font.size = Pt(18)
    line2.add_run("\t\t\t")
    l2b = line2.add_run(f"Language:{language or ''}")
    l2b.bold = True
    l2b.font.size = Pt(18)


def add_table_2_5(document: Document, class1_room: Optional[int], class2_room: Optional[int]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_time, hdr_sched = table.rows[0].cells
    p1 = hdr_time.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("TIME")
    r1.bold = True
    r1.font.size = Pt(14)
    p2 = hdr_sched.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SCHEDULE")
    r2.bold = True
    r2.font.size = Pt(14)
    for time_str, sched in SCHEDULE_ROWS_2_5:
        row = table.add_row()
        c1, c2 = row.cells
        p_time = c1.paragraphs[0]
        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = p_time.add_run(time_str)
        rt.bold = True
        rt.font.size = Pt(12)
        text = sched
        if "Reading/Writing-Class 1" in text and class1_room is not None:
            text = text.replace("Class 1", f"Class 1 (Room {class1_room})")
        if ("Math-Class 2" in text or "Science-Class 2" in text) and class2_room is not None:
            text = text.replace("Class 2", f"Class 2 (Room {class2_room})")
        p_sched = c2.paragraphs[0]
        p_sched.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rs = p_sched.add_run(text)
        rs.font.size = Pt(12)


def generate_2_5(frames: Dict[str, pd.DataFrame], outdir: Path) -> None:
    for sheet in ["2nd", "3rd", "4th", "5th"]:
        if sheet not in frames:
            console.print(f"[yellow]Skipping missing sheet:[/] {sheet}")
            continue
        df = frames[sheet]
        name_col = next((c for c in df.columns if isinstance(c, str) and c.lower().startswith("homeroom ")), None)
        lang_col = "Language" if "Language" in df.columns else None
        class1_col = "Class 1" if "Class 1" in df.columns else None
        class2_col = "Class 2" if "Class 2" in df.columns else None
        homeroom_num = None
        if isinstance(name_col, str):
            try:
                homeroom_num = int(name_col.split(" ", 1)[1])
            except Exception:
                pass
        document = Document()
        for _, row in df.iterrows():
            name = str(row.get(name_col, "")).strip()
            if not name:
                continue
            def to_int(v):
                try:
                    return int(v) if pd.notna(v) else None
                except Exception:
                    return None
            class1_room = to_int(row.get(class1_col)) if class1_col else None
            class2_room = to_int(row.get(class2_col)) if class2_col else None
            language = str(row.get(lang_col, "")).strip() if lang_col else None
            add_header(document, name, homeroom_num, sheet, language)
            add_table_2_5(document, class1_room, class2_room)
            document.add_page_break()
        out_path = outdir / (sheet.replace("/", "-").replace(" ", "_") + "_schedule.docx")
        document.save(out_path)
        console.print(f"[green]Wrote[/] {out_path}")


# ----------------------- 6th–8th schedules ----------------------- #

SCHEDULE_ROWS_6_8: List[tuple[str, str]] = [
    ("8:00-8:10", "Homeroom/Breakfast/Calm Classroom"),
    ("8:10-8:20", "*Second Step on Monday and Friday\n*MTSS on Tuesday, Wednesday, Thursday"),
    ("8:20-9:30", "Class 1 Science"),
    ("9:30-10:40", "Class 2 Reading/Writing"),
    ("10:40-11:10", "Class 3 Math"),
    ("11:10-12:10", "Student Enrichment Classes (See Below.)"),
    ("12:10-12:55", "Class 3  Math"),
    ("12:55-1:20", "Lunch-Lunchroom C"),
    ("1:20-1:50", "Recess-Indoor/Lunchroom A"),
    ("1:50-2:55", "Social Studies-All Homerooms"),
    ("2:55-3:00", "Homeroom/Dismissal"),
]


def add_header_6_8(document: Document, name: str, homeroom: Optional[int], grade: str, language: Optional[str]) -> None:
    line1 = document.add_paragraph()
    left = line1.add_run(f"Student:{name}")
    left.bold = True
    left.font.size = Pt(18)
    line1.add_run("\t\t\t")
    right = line1.add_run(f"{grade} Grade")
    right.bold = True
    right.font.size = Pt(18)
    line2 = document.add_paragraph()
    l2a = line2.add_run(f"Homeroom: {homeroom if homeroom is not None else ''}")
    l2a.bold = True
    l2a.font.size = Pt(18)
    line2.add_run("\t\t\t")
    l2b = line2.add_run(f"Language: {language or ''}")
    l2b.bold = True
    l2b.font.size = Pt(18)


def add_table_6_8(document: Document, c1: Optional[int], c2: Optional[int], c3: Optional[int]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_time, hdr_sched = table.rows[0].cells
    p1 = hdr_time.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("TIME")
    r1.bold = True
    r1.font.size = Pt(14)
    p2 = hdr_sched.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SCHEDULE")
    r2.bold = True
    r2.font.size = Pt(14)
    for time_str, sched in SCHEDULE_ROWS_6_8:
        row = table.add_row()
        c_time, c_sched = row.cells
        p_time = c_time.paragraphs[0]
        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = p_time.add_run(time_str)
        rt.bold = True
        rt.font.size = Pt(12)
        text = sched
        if "Class 1" in text and c1 is not None:
            text = text.replace("Class 1", f"Class 1 (Room {c1})")
        if "Class 2" in text and c2 is not None:
            text = text.replace("Class 2", f"Class 2 (Room {c2})")
        if "Class 3" in text and c3 is not None:
            text = text.replace("Class 3", f"Class 3 (Room {c3})")
        p_sched = c_sched.paragraphs[0]
        p_sched.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rs = p_sched.add_run(text)
        rs.font.size = Pt(12)


def generate_6_8(frames: Dict[str, pd.DataFrame], outdir: Path) -> None:
    for sheet in ["6th", "7th", "8th"]:
        if sheet not in frames:
            console.print(f"[yellow]Skipping missing sheet:[/] {sheet}")
            continue
        df = frames[sheet]
        name_col = next((c for c in df.columns if isinstance(c, str) and c.lower().startswith("homeroom ")), None)
        lang_col = "Language" if "Language" in df.columns else None
        c1 = next((c for c in df.columns if isinstance(c, str) and c.lower().startswith("class 1")), None)
        c2 = next((c for c in df.columns if isinstance(c, str) and c.lower().startswith("class 2")), None)
        c3 = next((c for c in df.columns if isinstance(c, str) and c.lower().startswith("class 3")), None)
        homeroom_num = None
        if isinstance(name_col, str):
            try:
                homeroom_num = int(name_col.split(" ", 1)[1])
            except Exception:
                pass
        document = Document()
        for _, row in df.iterrows():
            name = str(row.get(name_col, "")).strip()
            if not name:
                continue
            def to_int(v):
                try:
                    return int(v) if pd.notna(v) else None
                except Exception:
                    return None
            c1_room = to_int(row.get(c1)) if c1 else None
            c2_room = to_int(row.get(c2)) if c2 else None
            c3_room = to_int(row.get(c3)) if c3 else None
            language = str(row.get(lang_col, "")).strip() if lang_col else None
            add_header_6_8(document, name, homeroom_num, sheet, language)
            add_table_6_8(document, c1_room, c2_room, c3_room)
            document.add_page_break()
        out_path = outdir / (sheet.replace("/", "-").replace(" ", "_") + "_schedule.docx")
        document.save(out_path)
        console.print(f"[green]Wrote[/] {out_path}")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    workbook_path = Path(WORKBOOK_NAME)
    teacher_csv_path = Path(TEACHER_CSV_NAME)
    frames = load_workbook(workbook_path)
    teacher_map = load_teacher_mapping(teacher_csv_path)
    generate_cards(frames, teacher_map, OUTPUT_DIR)
    generate_2_5(frames, OUTPUT_DIR)
    generate_6_8(frames, OUTPUT_DIR)


if __name__ == "__main__":
    main()


